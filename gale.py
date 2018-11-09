from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from collections import OrderedDict
import csv
import re


global DATA


article 	    = 0
bold 		    = 1
keyword 	    = 2
seeword 	    = 3
books			= 4
periodicals 	= 5
organisations 	= 6
keytable	    = 7


ALL = OrderedDict()
DATA = OrderedDict()

document = Document('mental1.docx')
tables = document.tables

def iter_block_items(parent):
    # Reads and stores text in parent_elm
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    # Reads and stores tables in parent_elm
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def maincode():
    global document
    global tables
    global DATA
    temp_bold = []
    head = ''
    letter = ''
    a=1
    Books = 1
    Periodicals=1
    Organizations=1
    # define = 0
    file = [[],[],[], [],[],[], [],[],[]]
    
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            for run in block.runs:
                
                # Letter - A,B
                if(run.font.size in [584200, 577850, 685800, 571500, 565150]):
                    # print(run.text)
                    letter = run.text
                    ALL[letter]=DATA
                    
                 
                # Word Heading
                elif(run.font.size in [184150, 177800, 190500, 171450]):
                	Books=0
                	Periodicals=0
                	Organizations=0
                	a = 1
                	# define = 0
                	wd = run.text.lower()
                	if(wd in['see','acupressure']):
                		pass
                	else:
                		DATA[head.lower()] = file
                		head = run.text	
                		file = [[],[],[], [],[],[], [],[],[]]

                # elif(run.font.size in [139700]):
                # 	if(run.text.lower()=='definition'):
                # 		define = 1
                # 	else:
                # 		define = 0
                    
                
                # Bold
                elif(run.font.size == 120650):
                	file[bold].append(run.text)
                	# print(run.text)
                     
                # Seeword
                elif(run.font.size == 146050):
                    file[seeword].append(run.text)           
                    # print(run.text)          

                # Article
                elif(run.font.size == None):
         
                    # See also keyword
                    if(str(run.text).startswith('See also')):
                        a = 0

                    # if(define == 1):

                    # 	file[definition].append(run.text)
                    	
                    if a == 0:
                        # print(run.text)
                        file[keyword].append(run.text)
                    else:
                    	file[article].append(run.text)
        
                # References          
                elif(run.font.size == 101600):
                	define = 0
                	Books = 0
                	Periodicals = 0
                	Organizations = 0
                	if(str(run.text)=='BOOKS'):
                		Books = 1
                	elif(str(run.text)== 'PERIODICALS'):
                		Periodicals = 1
                	elif(str(run.text)== 'ORGANIZATIONS'):
                		Organizations = 1
                                
                # Author,Book name
                elif(run.font.size==114300):
                	if(Books == 1):
                		#print("\t\t\tBOOKS")
                		file[books].append(run.text)
                	elif(Periodicals == 1):
                		file[periodicals].append(run.text)
                	elif(Organizations == 1):
                		file[organisations].append(run.text)
                    # print(run.text)
                    
                # Unrequired font sizes  
                elif(run.font.size in [139700, 107950, 133350, 158750, 82550, 152400, 165100, 95250, ' '] ):
                	pass
                
                # Buffer
                else:
                	print(run.font.size)
                	print (run.text+"\t\t\tLack of a font size  ")

        elif isinstance(block, Table):
    
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            file[keytable].append(paragraph.text)
                            # print(paragraph.text)                      
                    break
                break
            tables = tables[1:]
            
    DATA[head]=file
    ALL[letter]=DATA
    
def clean():
    global DATA
    for word in DATA:
        for itr in range(0,9):
            
            cur = DATA[word][itr]
            temp = []
            for val in cur:
                val = val.strip('\n').strip(' ').strip(';').strip('<').strip('>')
                if not val=='':
                    temp.append(val)

            DATA[word][itr]=temp


    for key in DATA.keys():  
        DATA[key][article] = (' '.join(DATA[key][article]))
        DATA[key][keyword] = (' '.join(DATA[key][keyword]))
        DATA[key][keytable] = DATA[key][keytable][1:]

        if "also" in DATA[key][keyword].split():
            DATA[key][keyword] = DATA[key][keyword].split('also')[1].strip()

    for word in DATA:
        for itr in range(0,9):
            DATA[word][itr] = re.sub('\[|\]','', str(DATA[word][itr]))
            
    
def write():
    
    with open('mental1.csv', 'w') as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(["WORD", "ARTICLE", "BOLD", "KEYWORD", "SEE WORD", "BOOKS", "PERIODICALS", "ORGANISATIONS", "KEYTABLES"])
        for key, value in DATA.items(): 
            writer.writerow([key]+value)

if __name__=="__main__":

    maincode()
    clean()
    write()
    print("done")


